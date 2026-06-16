import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, font, ttk
import json
import os
import sys
import webbrowser
import re
import threading
import time
from langchain_text_splitters import RecursiveCharacterTextSplitter

# --- ПРОВЕРКА БИБЛИОТЕК ---
try:
    import docx
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_INPUT_AVAILABLE = True
except ImportError:
    PDF_INPUT_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import cm
    PDF_OUTPUT_AVAILABLE = True
except ImportError:
    PDF_OUTPUT_AVAILABLE = False

APP_NAME = "Чанкер РА"
AUTHOR_NAME = "berezowskiy"
TELEGRAM_LINK = "https://t.me/berezowskiy"

# --- ЦВЕТОВАЯ ПАЛИТРА ---
COLOR_BG = "#f4f6f9"
COLOR_CARD = "#ffffff"
COLOR_PRIMARY = "#4CAF50"
COLOR_PRIMARY_HOVER = "#43a047"
COLOR_SECONDARY = "#2196F3"
COLOR_SECONDARY_HOVER = "#1e88e5"
COLOR_TEXT = "#2c3e50"
COLOR_TEXT_MUTED = "#7f8c8d"
COLOR_ACCENT_BG = "#fff3e0"
COLOR_ACCENT_TEXT = "#e65100"
COLOR_SUCCESS_BG = "#e8f5e9"
COLOR_SUCCESS_TEXT = "#2e7d32"
COLOR_LINK = "#2980b9"
COLOR_LINK_HOVER = "#c0392b"

# --- МАППЕРЫ ---
TEMPLATE_NAMES = {
    "universal": "🏆 Универсальный (LangChain/Pinecone)",
    "speech_sense": "🤖 Speech Sense (Строгий парсинг)",
    "flat": "📄 Плоский текст (Ключевые слова)"
}
TEMPLATE_DESCS = {
    "universal": "Связный текст + мета-данные. Идеально для векторного поиска.",
    "speech_sense": "Конструкция 'Текст — это Список'. Для систем со строгим синтаксисом.",
    "flat": "Просто текст. Минимум структуры, максимум ключевых слов."
}
PRESET_NAMES = {
    "auto": "🤖 Авто-подбор (Рекомендуется)",
    "facts": "⚡ Точный поиск (FAQ, Факты)",
    "balanced": "⚖️ Универсальный баланс",
    "context": "📚 Глубокий контекст (Статьи)"
}
PRESET_DESCS = {
    "auto": "Программа сама проанализирует длину текста и подберет размер.",
    "facts": "Мелкие чанки (~400 симв). Для коротких ответов.",
    "balanced": "Средние чанки (~800 симв). Золотая середина.",
    "context": "Крупные чанки (~1200+ симв). Для сложных документов."
}

def get_resource_path(filename):
    # Для упакованного приложения (.app)
    if getattr(sys, 'frozen', False):
        if sys.platform == "darwin":  # macOS
            base_path = os.path.join(os.path.dirname(sys.executable), '..', 'Resources')
            path = os.path.join(base_path, filename)
            if os.path.exists(path):
                return path
        else:  # Windows
            path = os.path.join(sys._MEIPASS, filename)
            if os.path.exists(path):
                return path
    
    # Для разработки
    path = os.path.join(os.path.abspath("."), filename)
    if os.path.exists(path):
        return path
    
    return None

class ModernButton(tk.Button):
    def __init__(self, master, text, command, bg_color, hover_color, fg="white", pady_val=6, **kwargs):
        if 'font' not in kwargs:
            kwargs['font'] = ("Segoe UI", 11, "bold")
        super().__init__(master, text=text, command=command, bg=bg_color, fg=fg,
                         relief="flat", cursor="hand2",
                         pady=pady_val, activebackground=bg_color, activeforeground=fg, **kwargs)
        self.default_bg = bg_color
        self.hover_bg = hover_color
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)

    def on_enter(self, e):
        if self['state'] != 'disabled':
            self.config(bg=self.hover_bg)

    def on_leave(self, e):
        if self['state'] != 'disabled':
            self.config(bg=self.default_bg)

class SmartMappingDialog(tk.Toplevel):
    def __init__(self, parent, fields, data_sample, current_config):
        super().__init__(parent)
        self.title("🧠 Настройка полей")
        self.geometry("600x550") 
        self.configure(bg=COLOR_BG)
        self.result_config = None
        
        tk.Label(self, text="Подтвердите поля:", font=("Segoe UI", 11, "bold"), bg=COLOR_BG, pady=10).pack(fill="x", padx=15)
        
        self.vars = {}
        keywords = {
            "category_field": ["name", "category", "group", "type", "title", "label", "категория", "название"],
            "content_field": ["description", "text", "body", "content", "prompt", "desc", "описание", "текст"],
            "atomic_field": ["list", "tags", "items", "sub", "options", "categories", "second", "список", "теги"]
        }
        roles = [
            ("category_field", "🏷 Категория", "Группа записи"),
            ("content_field", "📄 Текст", "Основное описание"),
            ("atomic_field", "⚛ Списки", "Подкатегории")
        ]
        
        scroll_container = tk.Frame(self, bg=COLOR_BG)
        scroll_container.pack(fill="both", expand=True, padx=15, pady=5)
        self.canvas = tk.Canvas(scroll_container, bg=COLOR_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(scroll_container, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg=COLOR_CARD)
        
        self.scrollable_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar.set)
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        
        self.canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        frame = tk.Frame(self.scrollable_frame, bg=COLOR_CARD, pady=10, padx=15, relief="flat", bd=1)
        frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        for i, (key, label, hint) in enumerate(roles):
            lbl_frame = tk.Frame(frame, bg=COLOR_CARD)
            lbl_frame.grid(row=i, column=0, sticky="ew", pady=8)
            tk.Label(lbl_frame, text=label, font=("Segoe UI", 10, "bold"), bg=COLOR_CARD, anchor="w").pack(anchor="w")
            tk.Label(lbl_frame, text=hint, font=("Segoe UI", 8), fg=COLOR_TEXT_MUTED, bg=COLOR_CARD, anchor="w").pack(anchor="w")
            
            cb = ttk.Combobox(lbl_frame, values=["-- Не выбрано --"] + fields, state="readonly", font=("Segoe UI", 9))
            cb.pack(fill="x", pady=(4,0))
            
            selected = "-- Не выбрано --"
            current_val = current_config.get(key)
            if current_val in fields: selected = current_val
            else:
                for kw in keywords[key]:
                    for f in fields:
                        if kw.lower() in f.lower(): selected = f; break
                    if selected != "-- Не выбрано --": break
            
            if key == "atomic_field" and "categorySecondLine" in fields: selected = "categorySecondLine"
            cb.set(selected)
            self.vars[key] = cb
            cb.bind("<<ComboboxSelected>>", lambda e: self.update_preview(data_sample))
            
        frame.columnconfigure(0, weight=1)
        
        preview_frame = tk.LabelFrame(self.scrollable_frame, text="👁️ Предпросмотр", font=("Segoe UI", 9, "bold"), bg=COLOR_SUCCESS_BG, fg=COLOR_SUCCESS_TEXT, padx=10, pady=10)
        preview_frame.pack(fill="x", padx=10, pady=10)
        self.preview_text = tk.Text(preview_frame, height=5, wrap="word", font=("Consolas", 8), bg="#fff", fg="#333", relief="flat", bd=0)
        self.preview_text.pack(fill="x")
        self.preview_text.insert("1.0", "Настройте поля...")
        self.preview_text.config(state="disabled")
        
        btn_frame = tk.Frame(self, bg=COLOR_BG, pady=10)
        btn_frame.pack(side="bottom", fill="x")
        ModernButton(btn_frame, text="Отмена", command=self.destroy, bg_color="#95a5a6", hover_color="#7f8c8d", pady_val=6).pack(side="left", padx=10)
        ModernButton(btn_frame, text="✅ Принять", command=self.on_apply, bg_color=COLOR_PRIMARY, hover_color=COLOR_PRIMARY_HOVER, pady_val=6).pack(side="left", padx=10)
        
        self.data_sample = data_sample
        self.update_preview(data_sample)
        self.transient(parent)
        self.grab_set()
        self.after(100, lambda: self.canvas.yview_moveto(1.0))
        parent.wait_window(self)
        self.canvas.unbind_all("<MouseWheel>")

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    def update_preview(self, data):
        if not data or len(data) == 0: return
        cat_key = self.vars["category_field"].get()
        content_key = self.vars["content_field"].get()
        atomic_key = self.vars["atomic_field"].get()
        item = data[0]
        cat_val = item.get(cat_key, "???") if cat_key != "-- Не выбрано --" else "General"
        
        preview_str = ""
        if atomic_key != "-- Не выбрано --" and atomic_key in item and isinstance(item[atomic_key], list) and item[atomic_key]:
            sub_val = str(item[atomic_key][0]) if item[atomic_key] else "..."
            base_text = str(item.get(content_key, ""))[:60] if content_key != "-- Не выбрано --" else ""
            if len(str(item.get(content_key, ""))) > 60: base_text += "..."
            preview_str = f'Атомизация:\n{{ "cat": "{cat_val}", "text": "{base_text} - это \\"{sub_val}\\"" }}'
        else:
            content_val = str(item.get(content_key, ""))[:80] if content_key != "-- Не выбрано --" else "Нет текста"
            if len(str(item.get(content_key, ""))) > 80: content_val += "..."
            preview_str = f'Обычный:\n{{ "cat": "{cat_val}", "text": "{content_val}" }}'
            
        self.preview_text.config(state="normal")
        self.preview_text.delete("1.0", "end")
        self.preview_text.insert("1.0", preview_str)
        self.preview_text.config(state="disabled")

    def on_apply(self):
        config = {key: var.get() for key, var in self.vars.items()}
        final_config = {k: v for k, v in config.items() if v != "-- Не выбрано --"}
        self.result_config = final_config
        self.destroy()

class RAGChunkerApp:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_NAME)
        self.root.geometry("1000x800")
        self.root.configure(bg=COLOR_BG)
        self.root.resizable(True, True)

        self.processing = False
        self.input_file_path = tk.StringVar()
        self.mapping_config = {}
        self.available_fields = []
        self.raw_data_sample = None 
        self.is_json_mode = False
        
        self.chunk_size_var = tk.IntVar(value=800) 
        self.chunk_overlap_var = tk.IntVar(value=80)
        self.save_mode_var = tk.BooleanVar(value=False)
        self.output_template_var = tk.StringVar(value="universal")
        
        self.presets = {
            "auto": {"size": 800, "overlap": 80},
            "facts": {"size": 400, "overlap": 40},
            "balanced": {"size": 800, "overlap": 80},
            "context": {"size": 1200, "overlap": 150}
        }
        self.current_preset_var = tk.StringVar(value="auto")

        self.result_data = None
        self.source_filename = ""
        self.pdf_font_name = None

        self.font_title = font.Font(family="Segoe UI", size=12, weight="bold")
        self.font_label = font.Font(family="Segoe UI", size=9)
        self.font_log = font.Font(family="Consolas", size=8)
        self.font_footer = font.Font(family="Segoe UI", size=9, slant="italic")

        self.create_widgets()
        self.pdf_font_name = self._register_pdf_font()

    def _register_pdf_font(self):
        if not PDF_OUTPUT_AVAILABLE: 
            return None
        
        try:
            # Используем системный шрифт macOS
            self.log("Используется системный шрифт Helvetica", "info")
            return 'Helvetica'
        except Exception as e:
            self.log(f"Ошибка при выборе шрифта: {e}", "error")
            return 'Helvetica'  # fallback

    def create_widgets(self):
        self.main_canvas = tk.Canvas(self.root, bg=COLOR_BG, highlightthickness=0)
        self.main_scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=self.main_canvas.yview)
        self.scrollable_content = tk.Frame(self.main_canvas, bg=COLOR_BG)

        self.scrollable_content.bind("<Configure>", lambda e: self.main_canvas.configure(scrollregion=self.main_canvas.bbox("all")))
        self.main_canvas.create_window((0, 0), window=self.scrollable_content, anchor="nw", width=1000)
        
        self.root.bind("<Configure>", lambda e: self.main_canvas.itemconfig(self.main_canvas.find_withtag("all")[0], width=max(800, e.width-20)) if e.widget == self.root else None)
        self.main_canvas.bind_all("<MouseWheel>", lambda event: self.main_canvas.yview_scroll(int(-1*(event.delta/120)), "units"))

        self.main_canvas.pack(side="left", fill="both", expand=True)
        self.main_scrollbar.pack(side="right", fill="y")

        content = self.scrollable_content

        header_frame = tk.Frame(content, bg=COLOR_CARD, pady=15, padx=20)
        header_frame.pack(fill="x", pady=(20, 10))
        
        top_row = tk.Frame(header_frame, bg=COLOR_CARD)
        top_row.pack(fill="x")
        
        tk.Label(top_row, text="🐱", font=("Segoe UI Emoji", 35), bg=COLOR_CARD).pack(side="left", padx=10)
        title_box = tk.Frame(top_row, bg=COLOR_CARD)
        title_box.pack(side="left", fill="x", expand=True)
        tk.Label(title_box, text=APP_NAME, font=("Segoe UI", 20, "bold"), bg=COLOR_CARD, fg=COLOR_TEXT).pack(anchor="w")
        tk.Label(title_box, text="Профессиональный адаптер баз знаний с умным чанкингом", font=("Segoe UI", 10), bg=COLOR_CARD, fg=COLOR_TEXT_MUTED).pack(anchor="w")

        log_frame = tk.Frame(header_frame, bg="#263238", height=100)
        log_frame.pack(fill="x", pady=(15, 0))
        self.log_text = scrolledtext.ScrolledText(log_frame, height=4, state='disabled', font=self.font_log, bg="#263238", fg="#eceff1", relief="flat", bd=0)
        self.log_text.pack(fill="both", expand=True)
        self.log_text.tag_config("info", foreground="#4fc3f7")
        self.log_text.tag_config("success", foreground="#81c784")
        self.log_text.tag_config("error", foreground="#e57373")
        self.log_text.tag_config("warning", foreground="#f1c40f")
        self.log("Готов к работе. Загрузите файл.", "info")

        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(header_frame, variable=self.progress_var, maximum=100, mode='indeterminate')
        self.progress_bar.pack(fill="x", padx=20, pady=(0, 10))
        self.progress_bar.pack_forget() 

        step1 = tk.LabelFrame(content, text="📂 Шаг 1: Файл", font=self.font_title, padx=15, pady=15, bg=COLOR_CARD, fg=COLOR_TEXT, relief="flat", bd=1)
        step1.pack(fill="x", padx=20, pady=10)
        f1 = tk.Frame(step1, bg=COLOR_CARD)
        f1.pack(fill="x")
        tk.Entry(f1, textvariable=self.input_file_path, font=self.font_label, bg="#f8f9fa", relief="flat", bd=1, fg=COLOR_TEXT) \
          .pack(side="left", fill="x", expand=True, padx=(0, 15))
        ModernButton(f1, text="Обзор файлов", command=self.browse_input, bg_color="#95a5a6", hover_color="#7f8c8d", pady_val=8) \
          .pack(side="left")

        step2 = tk.LabelFrame(content, text="⚙️ Шаг 2: Настройки", font=self.font_title, padx=15, pady=15, bg=COLOR_CARD, fg=COLOR_TEXT, relief="flat", bd=1)
        step2.pack(fill="x", padx=20, pady=10)
        
        row_top = tk.Frame(step2, bg=COLOR_CARD)
        row_top.pack(fill="x", pady=5)
        self.btn_map = ModernButton(row_top, text="🧠 Настроить поля (Авто)", command=self.open_mapping_dialog, bg_color=COLOR_ACCENT_BG, hover_color="#ffe0b2", fg=COLOR_ACCENT_TEXT, pady_val=8, state="disabled")
        self.btn_map.pack(side="left", padx=5)
        self.lbl_map_status = tk.Label(row_top, text="Загрузите файл", font=("Segoe UI", 9), bg=COLOR_CARD, fg="#e74c3c")
        self.lbl_map_status.pack(side="left", padx=15)
        
        tk.Label(row_top, text="Шаблон:", font=self.font_label, bg=COLOR_CARD).pack(side="left", padx=(30, 10))
        template_values = list(TEMPLATE_NAMES.values())
        self.template_cb = ttk.Combobox(row_top, values=template_values, state="readonly", font=self.font_label, width=30)
        self.template_cb.pack(side="left")
        self.template_cb.bind("<<ComboboxSelected>>", self.on_template_change_beautiful)
        self.template_cb.set(TEMPLATE_NAMES[self.output_template_var.get()])

        self.lbl_template_desc = tk.Label(step2, text=TEMPLATE_DESCS["universal"], font=("Segoe UI", 8), fg=COLOR_TEXT_MUTED, bg=COLOR_CARD, wraplength=800, justify="left")
        self.lbl_template_desc.pack(anchor="w", padx=20, pady=(5, 10))

        row_bot = tk.Frame(step2, bg=COLOR_CARD)
        row_bot.pack(fill="x", pady=5)
        tk.Label(row_bot, text="Режим чанкинга:", font=self.font_label, bg=COLOR_CARD).pack(side="left", padx=5)
        preset_values = list(PRESET_NAMES.values())
        self.preset_cb = ttk.Combobox(row_bot, values=preset_values, state="readonly", font=self.font_label, width=25)
        self.preset_cb.pack(side="left", padx=10)
        self.preset_cb.bind("<<ComboboxSelected>>", self.on_preset_change_beautiful)
        self.preset_cb.set(PRESET_NAMES[self.current_preset_var.get()])
        
        self.lbl_preset_desc = tk.Label(step2, text=PRESET_DESCS["auto"], font=("Segoe UI", 8), fg=COLOR_TEXT_MUTED, bg=COLOR_CARD, wraplength=800, justify="left")
        self.lbl_preset_desc.pack(anchor="w", padx=20, pady=(5, 10))
        
        manual_frame = tk.Frame(step2, bg=COLOR_CARD)
        manual_frame.pack(fill="x", padx=20, pady=5)
        tk.Label(manual_frame, text="Размер:", font=self.font_label, bg=COLOR_CARD).pack(side="left", padx=(0, 5))
        self.sp_size = tk.Spinbox(manual_frame, from_=100, to=5000, textvariable=self.chunk_size_var, width=6, font=self.font_label, bd=1, relief="flat", bg="#f8f9fa")
        self.sp_size.pack(side="left")
        tk.Label(manual_frame, text="Перекрытие:", font=self.font_label, bg=COLOR_CARD).pack(side="left", padx=(20, 5))
        self.sp_overlap = tk.Spinbox(manual_frame, from_=0, to=1000, textvariable=self.chunk_overlap_var, width=6, font=self.font_label, bd=1, relief="flat", bg="#f8f9fa")
        self.sp_overlap.pack(side="left")
        btn_auto_tune = ModernButton(manual_frame, text="🔍 Авто-подбор", command=self.auto_tune_settings, bg_color=COLOR_SECONDARY, hover_color=COLOR_SECONDARY_HOVER, fg="white", pady_val=5)
        btn_auto_tune.pack(side="right", padx=5)

        step3 = tk.Frame(content, bg=COLOR_CARD, pady=15, relief="flat", bd=1)
        step3.pack(fill="x", padx=20, pady=10)
        tk.Label(step3, text="📥 Шаг 3: Сохранение", font=self.font_title, bg=COLOR_CARD, fg=COLOR_TEXT).pack(anchor="w", padx=20)
        
        cards_frame = tk.Frame(step3, bg=COLOR_CARD)
        cards_frame.pack(fill="x", padx=20, pady=10)
        
        card_read = tk.Frame(cards_frame, bg=COLOR_SUCCESS_BG, relief="flat", bd=1, cursor="hand2")
        card_read.pack(side="left", fill="x", expand=True, padx=10)
        card_read.bind("<Button-1>", lambda e: self.save_mode_var.set(False))
        tk.Radiobutton(card_read, variable=self.save_mode_var, value=False, bg=COLOR_SUCCESS_BG, activebackground=COLOR_SUCCESS_BG).pack(anchor="w", padx=15, pady=(10,0))
        tk.Label(card_read, text="👁️ Для чтения", font=("Segoe UI", 11, "bold"), bg=COLOR_SUCCESS_BG, fg=COLOR_SUCCESS_TEXT).pack(anchor="w", padx=15)
        tk.Label(card_read, text="Красивый документ (PDF/DOCX)", font=("Segoe UI", 9), bg=COLOR_SUCCESS_BG, fg="#333").pack(anchor="w", padx=15, pady=(0,10))

        card_rag = tk.Frame(cards_frame, bg=COLOR_ACCENT_BG, relief="flat", bd=1, cursor="hand2")
        card_rag.pack(side="left", fill="x", expand=True, padx=10)
        card_rag.bind("<Button-1>", lambda e: self.save_mode_var.set(True))
        tk.Radiobutton(card_rag, variable=self.save_mode_var, value=True, bg=COLOR_ACCENT_BG, activebackground=COLOR_ACCENT_BG).pack(anchor="w", padx=15, pady=(10,0))
        tk.Label(card_rag, text="🤖 Для Базы Знаний", font=("Segoe UI", 11, "bold"), bg=COLOR_ACCENT_BG, fg=COLOR_ACCENT_TEXT).pack(anchor="w", padx=15)
        tk.Label(card_rag, text="Чистый JSON для импорта", font=("Segoe UI", 9), bg=COLOR_ACCENT_BG, fg="#333").pack(anchor="w", padx=15, pady=(0,10))

        actions_frame = tk.Frame(content, bg=COLOR_BG, pady=20)
        actions_frame.pack(fill="x", padx=20, pady=10)
        self.btn_process = ModernButton(actions_frame, text="🚀 ЗАПУСТИТЬ ОБРАБОТКУ", command=self.process_data, bg_color=COLOR_PRIMARY, hover_color=COLOR_PRIMARY_HOVER, pady_val=15)
        self.btn_process.pack(fill="x", pady=(0, 15))
        self.btn_save = ModernButton(actions_frame, text="💾 СОХРАНИТЬ РЕЗУЛЬТАТ", command=self.save_output_menu, 
                                     bg_color=COLOR_SECONDARY, hover_color=COLOR_SECONDARY_HOVER, pady_val=15, state="disabled",
                                     font=("Segoe UI", 14, "bold"))
        self.btn_save.pack(fill="x")

        footer_frame = tk.Frame(content, bg=COLOR_BG, pady=30)
        footer_frame.pack(fill="x", padx=20, pady=10)
        tk.Label(footer_frame, text=f"© {AUTHOR_NAME}", font=self.font_footer, bg=COLOR_BG, fg=COLOR_TEXT_MUTED).pack(side="left")
        
        tg_label = tk.Label(footer_frame, text=f"✈️ Telegram: @{AUTHOR_NAME}", font=self.font_footer, bg=COLOR_BG, fg=COLOR_LINK, cursor="hand2")
        tg_label.pack(side="right")
        tg_label.bind("<Enter>", lambda e: tg_label.config(fg=COLOR_LINK_HOVER, font=(self.font_footer.family, 9, "italic", "underline")))
        tg_label.bind("<Leave>", lambda e: tg_label.config(fg=COLOR_LINK, font=(self.font_footer.family, 9, "italic")))
        tg_label.bind("<Button-1>", lambda e: webbrowser.open(TELEGRAM_LINK))

    def set_processing_state(self, is_processing):
        self.processing = is_processing
        state = 'disabled' if is_processing else 'normal'
        
        self.btn_process.config(state=state)
        self.btn_save.config(state=state)
        self.btn_map.config(state=state)
        self.sp_size.config(state=state)
        self.sp_overlap.config(state=state)
        self.template_cb.config(state=state if not is_processing else 'disabled')
        self.preset_cb.config(state=state if not is_processing else 'disabled')
        
        if is_processing:
            self.progress_bar.pack(fill="x", padx=20, pady=(0, 10))
            self.progress_bar.start(10)
            self.root.config(cursor="watch")
        else:
            self.progress_bar.stop()
            self.progress_var.set(0)
            self.progress_bar.pack_forget()
            self.root.config(cursor="")
        
        self.root.update_idletasks()

    def log(self, message, tag=None):
        if hasattr(self, 'log_text') and self.log_text:
            self.log_text.config(state='normal')
            self.log_text.insert(tk.END, f"[{APP_NAME}] {message}\n", tag)
            self.log_text.see(tk.END)
            self.log_text.config(state='disabled')
            self.root.update_idletasks()

    def on_template_change_beautiful(self, event):
        if self.processing: return
        beautiful_name = self.template_cb.get()
        tech_key = [k for k, v in TEMPLATE_NAMES.items() if v == beautiful_name][0]
        self.output_template_var.set(tech_key)
        self.lbl_template_desc.config(text=TEMPLATE_DESCS[tech_key])

    def on_preset_change_beautiful(self, event):
        if self.processing: return
        beautiful_name = self.preset_cb.get()
        tech_key = [k for k, v in PRESET_NAMES.items() if v == beautiful_name][0]
        self.current_preset_var.set(tech_key)
        self.lbl_preset_desc.config(text=PRESET_DESCS[tech_key])
        if tech_key == "auto": self.auto_tune_settings()
        else:
            data = self.presets[tech_key]
            self.chunk_size_var.set(data["size"])
            self.chunk_overlap_var.set(data["overlap"])

    def auto_tune_settings(self):
        if self.processing: return
        if not self.raw_data_sample and not self.is_json_mode:
            self.chunk_size_var.set(800); self.chunk_overlap_var.set(80)
            self.lbl_preset_desc.config(text=PRESET_DESCS["balanced"])
            self.current_preset_var.set("balanced"); self.preset_cb.set(PRESET_NAMES["balanced"])
            return
        if not self.raw_data_sample: return
        content_key = self.mapping_config.get("content_field", "description")
        lengths = []
        for item in self.raw_data_sample[:10]:
            if content_key in item:
                text = str(item[content_key])
                atomic_key = self.mapping_config.get("atomic_field")
                if atomic_key and atomic_key in item and isinstance(item[atomic_key], list) and item[atomic_key]:
                    avg_sub_len = sum(len(str(s)) for s in item[atomic_key]) / len(item[atomic_key])
                    lengths.append(len(text) + avg_sub_len + 20)
                else: lengths.append(len(text))
        if not lengths: lengths = [500]
        avg_len = sum(lengths) / len(lengths)
        selected_key = "balanced"
        if avg_len < 300: self.chunk_size_var.set(400); self.chunk_overlap_var.set(40); selected_key = "facts"
        elif avg_len > 1500: self.chunk_size_var.set(1000); self.chunk_overlap_var.set(150); selected_key = "context"
        else: self.chunk_size_var.set(800); self.chunk_overlap_var.set(80)
        self.lbl_preset_desc.config(text=PRESET_DESCS[selected_key])
        self.current_preset_var.set(selected_key)
        self.preset_cb.set(PRESET_NAMES[selected_key])

    def fast_aggressive_parse(self, text_content):
        """
        СУПЕР-БЫСТРЫЙ ПАРСЕР (БЕЗ REGEX В ЦИКЛЕ)
        Работает с индексами строк, а не создает новые объекты.
        """
        self.log("🚑 Запуск быстрого агрессивного восстановления (поток)...", "warning")
        
        start_idx = text_content.find('{')
        if start_idx == -1: return None
        end_idx = text_content.rfind('}')
        if end_idx == -1: return None
        
        # Грубая очистка от переносов для упрощения поиска
        block = text_content[start_idx:end_idx+1]
        block = block.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
        
        recovered = []
        # Разделяем по } {, но аккуратно
        parts = block.split('} {')
        
        for i, part in enumerate(parts):
            if not part.startswith('{'): part = '{' + part
            if not part.endswith('}'): part = part + '}'
            
            # Быстрая эвристика: ищем ключевые поля вручную
            obj = {}
            
            # Поиск name
            n_start = part.find('"name"')
            if n_start != -1:
                val_start = part.find(':', n_start)
                if val_start != -1:
                    q1 = part.find('"', val_start+1)
                    if q1 != -1:
                        q2 = part.find('"', q1+1)
                        # Ищем конец строки, учитывая возможные битые кавычки
                        # Упрощенно: берем до следующей запятой после кавычки
                        next_comma = part.find(',', q1+1)
                        next_brace = part.find('}', q1+1)
                        end_pos = min(next_comma if next_comma != -1 else 9999, next_brace if next_brace != -1 else 9999)
                        if end_pos != 9999:
                            # Проверяем, нет ли внутри незакрытых кавычек
                            segment = part[q1+1:end_pos]
                            # Если сегмент слишком длинный и имеет кавычки внутри - обрезаем по последней логичной
                            obj['name'] = segment.strip().rstrip('"').rstrip(',')
            
            # Поиск description
            d_start = part.find('"description"')
            if d_start != -1:
                val_start = part.find(':', d_start)
                if val_start != -1:
                    q1 = part.find('"', val_start+1)
                    if q1 != -1:
                        # Ищем конец описания (запятая или скобка, но не внутри кавычек)
                        # Для скорости берем до следующего ключа "categorySecondLine" или конца объекта
                        next_key = part.find('"categorySecondLine"', q1+1)
                        next_obj = part.find('},', q1+1) # Конец объекта в массиве
                        end_pos = 9999
                        if next_key != -1: end_pos = next_key
                        if next_obj != -1 and next_obj < end_pos: end_pos = next_obj
                        
                        if end_pos == 9999: end_pos = len(part) - 1
                        
                        segment = part[q1+1:end_pos]
                        # Чистим хвост
                        segment = segment.rstrip(',').rstrip('}').strip()
                        if segment.endswith('"'): segment = segment[:-1]
                        obj['description'] = segment.replace('\\"', '"')

            # Поиск списка
            l_start = part.find('"categorySecondLine"')
            if l_start != -1:
                val_start = part.find(':', l_start)
                if val_start != -1:
                    arr_start = part.find('[', val_start)
                    if arr_start != -1:
                        arr_end = part.find(']', arr_start)
                        if arr_end != -1:
                            arr_str = part[arr_start:arr_end+1]
                            try:
                                # Пытаемся распарсить только массив, он обычно целее
                                # Предварительно чистим разрывы внутри строк массива
                                arr_str_clean = re.sub(r'"\s*,\s*"', '","', arr_str) # Грубая фикс
                                obj['categorySecondLine'] = json.loads(arr_str_clean)
                            except:
                                obj['categorySecondLine'] = []

            if obj:
                recovered.append(obj)
            
            if i % 50 == 0:
                self.log(f"Обработано фрагментов: {i}", "info")

        if recovered:
            self.log(f"✅ Спасено объектов: {len(recovered)}", "success")
            return recovered
        return None

    def run_aggressive_thread(self, text_content):
        """Запуск парсера в отдельном потоке"""
        result = self.fast_aggressive_parse(text_content)
        
        # Возврат результата в главный поток
        def callback():
            self.set_processing_state(False)
            if result:
                self.raw_data_sample = result
                self.is_json_mode = True
                self.available_fields = list(result[0].keys()) if result else []
                self.btn_map.config(state="normal")
                self.lbl_map_status.config(text="Настройте поля (Режим Хирурга)", fg="#e67e22")
                self.log("Готово к настройке полей.", "success")
            else:
                self.log("❌ Не удалось восстановить данные даже в режиме Хирурга.", "error")
                self.lbl_map_status.config(text="Ошибка парсинга", fg="#c0392b")
        
        self.root.after(0, callback)

    def try_extract_json_from_text_robust(self, text_content):
        # Стандартные попытки (быстрые)
        start_brace = text_content.find('{')
        if start_brace == -1: return None
        
        clean_text = text_content[start_brace:].strip()
        try:
            data = json.loads(clean_text)
            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict): return data
            if isinstance(data, dict):
                for k, v in data.items():
                    if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict): return v
                return [data]
        except: pass

        # Попытка с удалением переносов
        candidate = re.sub(r'[\n\r\t]', '', clean_text[:clean_text.rfind('}')+1])
        try:
            data = json.loads(candidate)
            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict): return data
            if isinstance(data, dict):
                for k, v in data.items():
                    if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict): return v
                return [data]
        except: pass

        # Если не вышло - запускаем ХИРУРГА в потоке
        # Возвращаем None здесь, так как результат придет через callback
        self.log("Стандартные методы не сработали. Запуск режима Хирурга...", "warning")
        self.set_processing_state(True)
        thread = threading.Thread(target=self.run_aggressive_thread, args=(text_content,), daemon=True)
        thread.start()
        return None # Важно вернуть None, чтобы не блокировать UI

    def browse_input(self):
        if self.processing: return
        filetypes = [("Все файлы", "*.json *.txt *.docx *.pdf"), ("JSON", "*.json"), ("Text", "*.txt"), ("Word", "*.docx"), ("PDF", "*.pdf")]
        filename = filedialog.askopenfilename(filetypes=filetypes)
        if filename:
            self.input_file_path.set(filename)
            self.source_filename = os.path.basename(filename)
            ext = os.path.splitext(filename)[1].lower()
            self.log(f"Файл: {self.source_filename}", "info")
            self.is_json_mode = False; self.available_fields = []; self.raw_data_sample = None
            self.mapping_config = {}; self.btn_map.config(state="disabled")
            self.lbl_map_status.config(text="Анализ...", fg="#e74c3c")
            
            if ext == '.json': self._process_json_file(filename)
            elif ext == '.pdf': self._process_pdf_file(filename)
            elif ext in ['.txt', '.docx']: self._process_text_file(filename, ext)

    def _process_json_file(self, filename):
        try:
            with open(filename, 'r', encoding='utf-8') as f: data = json.load(f)
            if isinstance(data, dict):
                for k, v in data.items():
                    if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict): data = v; break
            if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                self.available_fields = list(data[0].keys())
                self.raw_data_sample = data; self.is_json_mode = True
                self.log(f"JSON найден. Поля: {', '.join(self.available_fields)}", "success")
                self.btn_map.config(state="normal")
                self.lbl_map_status.config(text="Настройте поля", fg="#e67e22")
        except Exception as e: self.log(f"Ошибка: {e}", "error")

    def _process_pdf_file(self, filename):
        if not PDF_INPUT_AVAILABLE:
            self.log("❌ Нет pdfplumber.", "error")
            return
        try:
            text_content = ""
            with pdfplumber.open(filename) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text: text_content += text + "\n"
            
            self.log("Анализ PDF...", "info")
            json_data = self.try_extract_json_from_text_robust(text_content)
            
            # Если вернулись данные сразу (стандартный путь)
            if json_data:
                self.log("✅ JSON найден!", "success")
                self.available_fields = list(json_data[0].keys())
                self.raw_data_sample = json_data
                self.is_json_mode = True
                self.btn_map.config(state="normal")
                self.lbl_map_status.config(text="Настройте поля", fg="#e67e22")
            # Если None - значит запущен поток, ждем callback
            elif not self.processing: 
                self.log("ℹ️ JSON не найден. Режим текста.", "warning")
                self.lbl_map_status.config(text="Простой текст", fg="#27ae60")
                self.mapping_config = {"content_field": "text"}
                self.auto_tune_settings()
        except Exception as e: self.log(f"Ошибка PDF: {e}", "error")

    def _process_text_file(self, filename, ext):
        try:
            text_content = ""
            if ext == '.docx':
                if not DOCX_AVAILABLE: self.log("❌ Нет python-docx", "error"); return
                doc = docx.Document(filename)
                text_content = "\n".join([p.text for p in doc.paragraphs])
            else:
                with open(filename, 'r', encoding='utf-8') as f: text_content = f.read()
            
            json_data = self.try_extract_json_from_text_robust(text_content)
            if json_data:
                self.log("✅ JSON найден!", "success")
                self.available_fields = list(json_data[0].keys())
                self.raw_data_sample = json_data; self.is_json_mode = True
                self.btn_map.config(state="normal")
                self.lbl_map_status.config(text="Настройте поля", fg="#e67e22")
            elif not self.processing:
                self.log("Режим простого текста", "warning")
                self.lbl_map_status.config(text="Простой текст", fg="#27ae60")
                self.mapping_config = {"content_field": "text"}
                self.auto_tune_settings()
        except Exception as e: self.log(f"Ошибка: {e}", "error")

    def open_mapping_dialog(self):
        if self.processing: return
        if not self.available_fields:
            messagebox.showwarning("Внимание", "Сначала загрузите файл.")
            return
        dialog = SmartMappingDialog(self.root, self.available_fields, self.raw_data_sample, self.mapping_config)
        if dialog.result_config:
            self.mapping_config = dialog.result_config
            self.log("Настройки применены", "success")
            self.lbl_map_status.config(text="Готово ✅", fg="#27ae60")
            self.auto_tune_settings()

    def normalize_and_chunk(self, data):
        chunks = []
        cat_field = self.mapping_config.get("category_field")
        content_field = self.mapping_config.get("content_field")
        atomic_field = self.mapping_config.get("atomic_field")
        template_type = self.output_template_var.get()
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size_var.get(), chunk_overlap=self.chunk_overlap_var.get(),
            length_function=len, separators=["\n\n", "\n", ". ", " ", ""]
        )
        total_items = len(data)
        
        for idx, item in enumerate(data):
            if idx % 50 == 0:
                self.log(f"Чанкинг: {idx}/{total_items}...", "info")
                self.root.update_idletasks()

            category = str(item.get(cat_field, "Unknown")) if cat_field else "General"
            base_text = str(item.get(content_field, "")).strip() if content_field else ""
            sub_list = []
            if atomic_field and atomic_field in item and isinstance(item[atomic_field], list): sub_list = item[atomic_field]
            items_to_process = []
            if not sub_list:
                items_to_process.append({"category": category, "content": base_text})
            else:
                if template_type == "speech_sense":
                    for sub_item in sub_list:
                        sub_str = str(sub_item).strip()
                        if not sub_str: continue
                        txt = f"{base_text} - это \"{sub_str}\"" if base_text else f"Элемент: \"{sub_str}\""
                        items_to_process.append({"category": category, "content": txt})
                elif template_type == "universal":
                    for sub_item in sub_list:
                        sub_str = str(sub_item).strip()
                        if not sub_str: continue
                        txt = f"В категории '{category}': {base_text}. Аспект: {sub_str}." if base_text else f"Аспект '{category}': {sub_str}."
                        items_to_process.append({"category": category, "content": txt})
                elif template_type == "flat":
                    list_str = ", ".join([str(s) for s in sub_list])
                    txt = f"{base_text} {list_str}" if base_text else list_str
                    items_to_process.append({"category": category, "content": txt})
            for unit in items_to_process:
                if not unit["content"].strip(): continue
                split_texts = splitter.split_text(unit["content"])
                for i, chunk_text in enumerate(split_texts):
                    chunk_obj = {}
                    if template_type == "universal":
                        chunk_obj = {"content": chunk_text, "metadata": {"category": unit["category"], "source": self.source_filename, "chunk_id": f"{idx}_{i}"}}
                    elif template_type == "speech_sense":
                        chunk_obj = {"category": unit["category"], "content": chunk_text, "chunk_id": f"{idx}_{i}"}
                    else:
                        chunk_obj = {"text": chunk_text, "category": unit["category"]}
                    chunks.append(chunk_obj)
        return chunks

    def process_data(self):
        if self.processing: return
        input_path = self.input_file_path.get()
        if not input_path or not os.path.exists(input_path):
            messagebox.showerror("Ошибка", "Выберите файл!")
            return
        if self.is_json_mode and not self.mapping_config:
            messagebox.showwarning("Внимание", "Настройте поля!")
            return
        
        self.set_processing_state(True)
        
        try:
            self.log("Обработка...", "info")
            if self.is_json_mode: data = self.raw_data_sample
            else:
                ext = os.path.splitext(input_path)[1].lower()
                if ext == '.docx' and DOCX_AVAILABLE:
                    doc = docx.Document(input_path)
                    full_text = "\n".join([p.text for p in doc.paragraphs])
                elif ext == '.pdf' and PDF_INPUT_AVAILABLE:
                    with pdfplumber.open(input_path) as pdf:
                        full_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                else:
                    with open(input_path, 'r', encoding='utf-8') as f: full_text = f.read()
                data = [{"text": full_text}]
                self.mapping_config = {"content_field": "text", "category_field": None}
            
            self.result_data = self.normalize_and_chunk(data)
            self.log(f"Готово! Чанков: {len(self.result_data)}", "success")
            self.btn_save.config(state="normal")
        except Exception as e:
            self.log(f"Ошибка: {str(e)}", "error")
            messagebox.showerror("Ошибка обработки", str(e))
        finally:
            self.set_processing_state(False)

    def save_output_menu(self):
        if self.processing: return
        if not self.result_data: return
        filetypes = [("PDF Document", "*.pdf"), ("JSON", "*.json"), ("Text", "*.txt"), ("Word", "*.docx")]
        filename = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=filetypes)
        if not filename: return
        ext = os.path.splitext(filename)[1].lower()
        
        self.set_processing_state(True)
        self.log("Начало сохранения...", "info")
        
        try:
            if ext == '.json': self.save_json(filename)
            elif ext == '.txt': self.save_txt(filename)
            elif ext == '.docx': self.save_docx(filename)
            elif ext == '.pdf': self.save_pdf(filename)
            mode = "БЗ" if self.save_mode_var.get() else "Чтение"
            self.log(f"Сохранено ({mode}): {filename}", "success")
            messagebox.showinfo("Успех", f"Готово!\nРежим: {mode}\nФайл: {filename}")
        except Exception as e:
            self.log(f"Ошибка сохранения: {str(e)}", "error")
            messagebox.showerror("Ошибка сохранения", str(e))
        finally:
            self.set_processing_state(False)

    def save_json(self, path):
        with open(path, 'w', encoding='utf-8') as f: json.dump(self.result_data, f, ensure_ascii=False, indent=2)

    def save_txt(self, path):
        is_rag = self.save_mode_var.get()
        with open(path, 'w', encoding='utf-8') as f:
            if is_rag: json.dump(self.result_data, f, ensure_ascii=False, indent=2)
            else:
                for i, chunk in enumerate(self.result_data):
                    cat = chunk.get('category', chunk.get('metadata', {}).get('category', 'N/A'))
                    f.write(f"### ЧАНК #{i+1} | {cat}\n{chunk.get('content', chunk.get('text', ''))}\n\n")

    def save_docx(self, path):
        if not DOCX_AVAILABLE: raise ImportError("Нет python-docx")
        doc = docx.Document()
        is_rag = self.save_mode_var.get()
        if is_rag:
            p = doc.add_paragraph(json.dumps(self.result_data, ensure_ascii=False, indent=2))
            p.style.font.name = 'Consolas'; p.style.font.size = Pt(8)
        else:
            for i, chunk in enumerate(self.result_data):
                if i % 100 == 0:
                    self.log(f"Запись DOCX: {i}/{len(self.result_data)}...", "info")
                    self.root.update_idletasks()
                cat = chunk.get('category', chunk.get('metadata', {}).get('category', 'N/A'))
                h = doc.add_heading(f"Чанк #{i+1}: {cat}", level=1)
                h.style.font.color.rgb = docx.shared.RGBColor(33, 150, 243)
                doc.add_paragraph(chunk.get('content', chunk.get('text', ''))).style.font.size = Pt(11)
                if i < len(self.result_data) - 1: doc.add_paragraph("_" * 30)
        doc.save(path)

    def save_pdf(self, path):
        if not PDF_OUTPUT_AVAILABLE: raise ImportError("Нет reportlab")
        if not self.pdf_font_name: raise RuntimeError("Нет шрифта")
        
        self.log("Генерация PDF...", "info")
        self.root.update_idletasks()

        doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=1*cm, leftMargin=1*cm, topMargin=1*cm, bottomMargin=1*cm)
        story = []
        styles = getSampleStyleSheet()
        is_rag = self.save_mode_var.get()
        
        if is_rag:
            code_style = ParagraphStyle('CodeRU', parent=styles['Normal'], fontName=self.pdf_font_name, fontSize=7, leading=9, textColor=colors.black)
            json_str = json.dumps(self.result_data, ensure_ascii=False, indent=2)
            lines = json_str.split('\n')
            for i, line in enumerate(lines):
                if i % 500 == 0:
                    self.log(f"Запись PDF: {i}/{len(lines)}...", "info")
                    self.root.update_idletasks()
                safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_line, code_style))
        else:
            title_style = ParagraphStyle('TitleRU', parent=styles['Heading1'], fontName=self.pdf_font_name, fontSize=14, spaceAfter=10, textColor=colors.darkblue)
            normal_style = ParagraphStyle('NormalRU', parent=styles['Normal'], fontName=self.pdf_font_name, fontSize=8, leading=10)
            
            story.append(Paragraph("Отчет Чанкер РА", title_style))
            story.append(Paragraph(f"Чанков: {len(self.result_data)}", normal_style))
            story.append(Spacer(1, 10))
            
            data = [["ID", "Категория", "Текст"]]
            total_chunks = len(self.result_data)
            
            for i, chunk in enumerate(self.result_data):
                if i % 100 == 0:
                    self.log(f"Запись таблицы: {i}/{total_chunks}...", "info")
                    self.root.update_idletasks()
                
                content_safe = chunk['content'].replace('\n', '<br/>')
                p_content = Paragraph(content_safe, normal_style)
                cat_val = chunk.get('category', chunk.get('metadata', {}).get('category', 'N/A'))
                data.append([str(chunk.get('chunk_id', '-')), cat_val, p_content])
            
            table = Table(data, colWidths=[1.5*cm, 4*cm, 12.5*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), self.pdf_font_name),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTNAME', (0, 1), (-1, -1), self.pdf_font_name),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('WORDWRAP', (0, 0), (-1, -1), True),
                ('SPLITROW', (0, 0), (-1, -1), True),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.beige, colors.white]),
            ]))
            story.append(table)
        
        doc.build(story)
        self.log("PDF создан!", "success")

if __name__ == "__main__":
    root = tk.Tk()
    
    # Для Windows
    if sys.platform == "win32":
        icon_path = get_resource_path("cat_icon.ico")
        if icon_path and os.path.exists(icon_path):
            try: root.iconbitmap(icon_path)
            except: pass
    # Для macOS - иконка устанавливается через py2app
    
    app = RAGChunkerApp(root)
    root.mainloop()
